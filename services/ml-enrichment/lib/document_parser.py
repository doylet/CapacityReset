"""
Document Parser Utilities - Professional Document Processing

Handles parsing and text extraction from various professional document formats.
Supports PDF, DOC/DOCX, TXT, and web-based document sources for brand analysis.
"""

import re
import io
import logging
import requests
import tempfile
from typing import Optional, Dict, Any, Tuple
from pathlib import Path
from urllib.parse import urlparse
import mimetypes

# Import document processing libraries
try:
    import PyPDF2
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    
try:
    import docx
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False


class DocumentParser:
    """
    Parses professional documents from various sources and formats.
    
    Supports:
    - PDF documents (CV, resume, portfolio)
    - Word documents (.doc/.docx)
    - Plain text files
    - LinkedIn profile URLs
    - Web-based document URLs
    """
    
    def __init__(self):
        """Initialize DocumentParser with logging."""
        self.logger = logging.getLogger(__name__)
        self.supported_formats = self._get_supported_formats()
    
    def _get_supported_formats(self) -> Dict[str, bool]:
        """Get available document format support based on installed libraries."""
        return {
            'pdf': PDF_AVAILABLE,
            'docx': DOCX_AVAILABLE,
            'html': BS4_AVAILABLE,
            'txt': True  # Always available
        }
    
    async def parse_document_from_url(self, document_url: str) -> Tuple[str, Dict[str, Any]]:
        """
        Parse document from URL and return text content with metadata.
        
        Args:
            document_url: URL to document or LinkedIn profile
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        self.logger.info(f"Parsing document from URL: {document_url}")
        
        # Determine document type from URL
        doc_type, source_type = self._analyze_document_url(document_url)
        
        if source_type == 'linkedin':
            return await self._parse_linkedin_profile(document_url)
        elif source_type == 'web_document':
            return await self._parse_web_document(document_url)
        elif source_type == 'file_url':
            return await self._parse_file_from_url(document_url, doc_type)
        else:
            raise ValueError(f"Unsupported document URL format: {document_url}")
    
    async def parse_document_from_file(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Parse document from local file path.
        
        Args:
            file_path: Local path to document file
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        self.logger.info(f"Parsing local document: {file_path}")
        
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Document file not found: {file_path}")
        
        # Determine file type from extension
        file_extension = path.suffix.lower()
        doc_type = self._extension_to_type(file_extension)
        
        # Parse based on document type
        if doc_type == 'pdf':
            return await self._parse_pdf_file(file_path)
        elif doc_type == 'docx':
            return await self._parse_docx_file(file_path)
        elif doc_type == 'txt':
            return await self._parse_text_file(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    async def parse_document_text(self, document_text: str, source_url: str = None) -> Tuple[str, Dict[str, Any]]:
        """
        Parse document from provided text content.
        
        Args:
            document_text: Raw document text content
            source_url: Optional source URL for metadata
            
        Returns:
            Tuple of (cleaned_text, metadata)
        """
        self.logger.info("Parsing provided document text")
        
        # Clean and normalize text
        cleaned_text = self._clean_document_text(document_text)
        
        # Extract metadata from text analysis
        metadata = {
            'source_type': 'direct_text',
            'source_url': source_url,
            'word_count': len(cleaned_text.split()),
            'char_count': len(cleaned_text),
            'detected_sections': self._detect_document_sections(cleaned_text),
            'processing_timestamp': self._get_current_timestamp()
        }
        
        return cleaned_text, metadata
    
    def _analyze_document_url(self, url: str) -> Tuple[str, str]:
        """Analyze URL to determine document type and source type."""
        url_lower = url.lower()
        
        # LinkedIn profile detection
        if 'linkedin.com' in url_lower and '/in/' in url_lower:
            return 'linkedin', 'linkedin'
        
        # Parse URL for file extension
        parsed = urlparse(url)
        path = parsed.path.lower()
        
        # File URL detection
        if path.endswith('.pdf'):
            return 'pdf', 'file_url'
        elif path.endswith(('.doc', '.docx')):
            return 'docx', 'file_url'
        elif path.endswith('.txt'):
            return 'txt', 'file_url'
        
        # Default to web document
        return 'html', 'web_document'
    
    def _extension_to_type(self, extension: str) -> str:
        """Convert file extension to document type."""
        extension_map = {
            '.pdf': 'pdf',
            '.doc': 'docx',
            '.docx': 'docx',
            '.txt': 'txt',
            '.html': 'html',
            '.htm': 'html'
        }
        return extension_map.get(extension, 'unknown')
    
    async def _parse_linkedin_profile(self, profile_url: str) -> Tuple[str, Dict[str, Any]]:
        """Parse LinkedIn profile URL for professional information."""
        
        # Note: This is a simplified implementation
        # In production, would need proper LinkedIn API integration
        # or specialized scraping with respect to ToS
        
        try:
            # For now, return placeholder text indicating LinkedIn profile
            # In full implementation, would extract:
            # - Professional headline
            # - Summary/About section  
            # - Experience descriptions
            # - Skills and endorsements
            
            extracted_text = f"LinkedIn Profile: {profile_url}\n\n"
            extracted_text += "Note: LinkedIn profile content extraction requires API access or specialized scraping. "
            extracted_text += "Please provide resume/CV content directly for detailed brand analysis."
            
            metadata = {
                'source_type': 'linkedin_profile',
                'source_url': profile_url,
                'extraction_method': 'placeholder',
                'processing_timestamp': self._get_current_timestamp(),
                'notes': 'Requires LinkedIn API or direct content provision'
            }
            
            return extracted_text, metadata
            
        except Exception as e:
            self.logger.error(f"Error parsing LinkedIn profile: {str(e)}")
            raise ValueError(f"Failed to parse LinkedIn profile: {profile_url}")
    
    async def _parse_web_document(self, document_url: str) -> Tuple[str, Dict[str, Any]]:
        """Parse web-based document URL."""
        
        if not BS4_AVAILABLE:
            raise ValueError("BeautifulSoup4 required for web document parsing")
        
        try:
            # Fetch document content
            headers = {
                'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36'
            }
            response = requests.get(document_url, headers=headers, timeout=30)
            response.raise_for_status()
            
            # Parse HTML content
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text content
            text_content = soup.get_text()
            
            # Clean and normalize
            cleaned_text = self._clean_document_text(text_content)
            
            metadata = {
                'source_type': 'web_document',
                'source_url': document_url,
                'content_type': response.headers.get('content-type', 'unknown'),
                'word_count': len(cleaned_text.split()),
                'processing_timestamp': self._get_current_timestamp()
            }
            
            return cleaned_text, metadata
            
        except Exception as e:
            self.logger.error(f"Error parsing web document: {str(e)}")
            raise ValueError(f"Failed to parse web document: {document_url}")
    
    async def _parse_file_from_url(self, file_url: str, doc_type: str) -> Tuple[str, Dict[str, Any]]:
        """Download and parse file from URL."""
        
        try:
            # Download file
            response = requests.get(file_url, timeout=30)
            response.raise_for_status()
            
            # Create temporary file
            with tempfile.NamedTemporaryFile(suffix=f'.{doc_type}', delete=False) as tmp_file:
                tmp_file.write(response.content)
                tmp_file_path = tmp_file.name
            
            # Parse the downloaded file
            if doc_type == 'pdf':
                text, metadata = await self._parse_pdf_file(tmp_file_path)
            elif doc_type == 'docx':
                text, metadata = await self._parse_docx_file(tmp_file_path)
            elif doc_type == 'txt':
                text, metadata = await self._parse_text_file(tmp_file_path)
            else:
                raise ValueError(f"Unsupported file type: {doc_type}")
            
            # Update metadata with URL source
            metadata['source_url'] = file_url
            metadata['download_method'] = 'http_request'
            
            # Cleanup temporary file
            Path(tmp_file_path).unlink(missing_ok=True)
            
            return text, metadata
            
        except Exception as e:
            self.logger.error(f"Error parsing file from URL: {str(e)}")
            raise ValueError(f"Failed to parse file from URL: {file_url}")
    
    async def _parse_pdf_file(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Parse PDF document file."""
        
        if not PDF_AVAILABLE:
            raise ValueError("PyPDF2 required for PDF parsing")
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                
                # Extract text from all pages
                text_content = ""
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n\n"
                
                # Clean and normalize
                cleaned_text = self._clean_document_text(text_content)
                
                metadata = {
                    'source_type': 'pdf_file',
                    'source_path': file_path,
                    'page_count': len(pdf_reader.pages),
                    'word_count': len(cleaned_text.split()),
                    'processing_timestamp': self._get_current_timestamp()
                }
                
                return cleaned_text, metadata
                
        except Exception as e:
            self.logger.error(f"Error parsing PDF file: {str(e)}")
            raise ValueError(f"Failed to parse PDF file: {file_path}")
    
    async def _parse_docx_file(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Parse Word document file."""
        
        if not DOCX_AVAILABLE:
            raise ValueError("python-docx required for DOCX parsing")
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Clean and normalize
            cleaned_text = self._clean_document_text(text_content)
            
            metadata = {
                'source_type': 'docx_file',
                'source_path': file_path,
                'paragraph_count': len(doc.paragraphs),
                'word_count': len(cleaned_text.split()),
                'processing_timestamp': self._get_current_timestamp()
            }
            
            return cleaned_text, metadata
            
        except Exception as e:
            self.logger.error(f"Error parsing DOCX file: {str(e)}")
            raise ValueError(f"Failed to parse DOCX file: {file_path}")
    
    async def _parse_text_file(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Parse plain text file."""
        
        try:
            # Try multiple encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            text_content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text_content = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if text_content is None:
                raise ValueError("Could not decode file with any supported encoding")
            
            # Clean and normalize
            cleaned_text = self._clean_document_text(text_content)
            
            metadata = {
                'source_type': 'text_file',
                'source_path': file_path,
                'word_count': len(cleaned_text.split()),
                'encoding_used': encoding,
                'processing_timestamp': self._get_current_timestamp()
            }
            
            return cleaned_text, metadata
            
        except Exception as e:
            self.logger.error(f"Error parsing text file: {str(e)}")
            raise ValueError(f"Failed to parse text file: {file_path}")
    
    def _clean_document_text(self, raw_text: str) -> str:
        """Clean and normalize document text for analysis."""
        
        if not raw_text:
            return ""
        
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', raw_text)
        
        # Remove excessive line breaks
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        
        # Clean up common document artifacts
        text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', text)  # Remove control characters
        text = re.sub(r'•\s*', '• ', text)  # Normalize bullet points
        text = re.sub(r'\s*-\s*', ' - ', text)  # Normalize dashes
        
        # Trim and return
        return text.strip()
    
    def _detect_document_sections(self, text: str) -> Dict[str, bool]:
        """Detect common professional document sections."""
        
        text_lower = text.lower()
        
        sections = {
            'contact_info': bool(re.search(r'(email|phone|linkedin|address)', text_lower)),
            'summary': bool(re.search(r'(summary|profile|objective)', text_lower)),
            'experience': bool(re.search(r'(experience|employment|work|position)', text_lower)),
            'education': bool(re.search(r'(education|degree|university|college)', text_lower)),
            'skills': bool(re.search(r'(skills|competencies|abilities)', text_lower)),
            'achievements': bool(re.search(r'(achievement|award|recognition)', text_lower)),
            'projects': bool(re.search(r'(project|portfolio)', text_lower))
        }
        
        return sections
    
    def _get_current_timestamp(self) -> str:
        """Get current timestamp string."""
        from datetime import datetime
        return datetime.utcnow().isoformat()
    
    def get_parser_capabilities(self) -> Dict[str, Any]:
        """Get information about parser capabilities and requirements."""
        
        return {
            'supported_formats': self.supported_formats,
            'requirements': {
                'pdf': 'pip install PyPDF2' if not PDF_AVAILABLE else 'Available',
                'docx': 'pip install python-docx' if not DOCX_AVAILABLE else 'Available',
                'html': 'pip install beautifulsoup4' if not BS4_AVAILABLE else 'Available',
                'txt': 'Built-in support'
            },
            'url_types': ['linkedin_profiles', 'web_documents', 'direct_file_urls'],
            'file_types': ['pdf', 'docx', 'txt', 'html']
        }